import os
import tempfile
from dotenv import load_dotenv
from telegram import Update
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackContext
from docx import Document
import fitz  # PyMuPDF
import textract
from openai import OpenAI
from langdetect import detect_langs
import tiktoken

# === Загрузка переменных окружения ===
load_dotenv()
TELEGRAM_TOKEN = os.getenv('TELEGRAM_TOKEN')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
ADMIN_ID = int(os.getenv('ADMIN_ID'))

client = OpenAI(api_key=OPENAI_API_KEY)
user_files = {}
updater = None

# === Подсчёт токенов ===
def count_tokens(text, model="gpt-4o"):
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))

# === Извлечение текста ===
def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    return "\n".join([page.get_text() for page in doc])

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_text_from_doc(file_path):
    return textract.process(file_path).decode('utf-8')

# === GPT-сравнение (всегда на русском) ===
def ask_openai_to_compare(text1, text2):
    prompt = f"""Сравни два технических документа. Построй подробный отчёт по следующим пунктам:

1. Совпадения — перечисли параметры, которые совпадают в обоих документах.
2. Различия — перечисли параметры, значения которых отличаются.
3. Таблица сравнения — оформи в виде таблицы:

| Параметр | Значение в документе 1 | Значение в документе 2 | Совпадает (Да/Нет) |

(если значение отсутствует — так и укажи)

4. Общий вывод — краткое описание основных различий и совпадений, без оценки качества документов.

Документ 1:
{text1}

Документ 2:
{text2}
"""

    prompt_tokens = count_tokens(prompt, "gpt-4o")
    max_total_tokens = 128000
    max_completion_tokens = 16000
    available_tokens = min(max_completion_tokens, max_total_tokens - prompt_tokens)

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=available_tokens
        )
        return response.choices[0].message.content
    except Exception as e:
        raise Exception(f"Ошибка OpenAI: {e}")

# === Сохранение в docx с таблицей ===
def save_docx_from_text(text):
    doc = Document()
    doc.add_heading("Сравнительный анализ документов", level=1)

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    table_lines = [line for line in lines if " | " in line]

    if table_lines:
        headers = table_lines[0].split(" | ")
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            table.cell(0, i).text = header.strip()

        for line in table_lines[1:]:
            cells = line.split(" | ")
            row = table.add_row().cells
            for i, cell in enumerate(cells):
                if i < len(row):
                    row[i].text = cell.strip()

        other_text = [line for line in lines if line not in table_lines]
        if other_text:
            doc.add_paragraph("\n".join(other_text))
    else:
        for line in lines:
            doc.add_paragraph(line)

    result_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
    doc.save(result_path)
    return result_path

# === Команды Telegram ===
def start(update: Update, context: CallbackContext):
    update.message.reply_text("Привет! Пришли два файла (PDF, DOCX, DOC или TXT). Сначала эталон, затем сравниваемый документ.")

def stop(update: Update, context: CallbackContext):
    global updater
    user_id = update.message.from_user.id

    if user_id == ADMIN_ID:
        update.message.reply_text("Бот завершает работу.")
        if updater:
            updater.stop()
            updater.is_idle = False
    else:
        update.message.reply_text("У вас нет прав для остановки бота.")

# === Обработка файлов ===
def handle_file(update: Update, context: CallbackContext):
    user_id = update.message.from_user.id
    if user_id not in user_files:
        user_files[user_id] = []

    file = update.message.document.get_file()
    file_name = update.message.document.file_name
    file_ext = os.path.splitext(file_name)[1].lower()
    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=file_ext).name
    file.download(custom_path=temp_path)

    try:
        if file_ext == '.pdf':
            text = extract_text_from_pdf(temp_path)
        elif file_ext == '.docx':
            text = extract_text_from_docx(temp_path)
        elif file_ext == '.doc':
            text = extract_text_from_doc(temp_path)
        elif file_ext == '.txt':
            text = extract_text_from_txt(temp_path)
        else:
            update.message.reply_text("Формат не поддерживается. Отправь PDF, DOCX, DOC или TXT.")
            os.remove(temp_path)
            return
    except Exception as e:
        update.message.reply_text(f"Ошибка при чтении файла: {e}")
        os.remove(temp_path)
        return

    user_files[user_id].append((temp_path, text))
    update.message.reply_text(f"Файл {len(user_files[user_id])} получен.")

    if len(user_files[user_id]) == 2:
        (_, text1), (_, text2) = user_files[user_id]
        update.message.reply_text("Анализирую документы...")

        try:
            if not text1.strip() or not text2.strip():
                update.message.reply_text("Один из документов пустой или не удалось извлечь текст. Проверь формат.")
                return

            result_text = ask_openai_to_compare(text1, text2)
        except Exception as e:
            update.message.reply_text(str(e))
            return

        update.message.reply_text("Анализ завершён. Отправляю результат в формате DOCX...")

        docx_path = save_docx_from_text(result_text)
        update.message.reply_document(open(docx_path, "rb"), filename="Сравнение.docx")

        for f, _ in user_files.get(user_id, []):
            os.remove(f)
        if os.path.exists(docx_path):
            os.remove(docx_path)

        user_files[user_id] = []

# === Запуск бота ===
def main():
    global updater
    updater = Updater(TELEGRAM_TOKEN, use_context=True)
    dp = updater.dispatcher

    dp.add_handler(CommandHandler("start", start))
    dp.add_handler(CommandHandler("stop", stop))
    dp.add_handler(MessageHandler(Filters.document, handle_file))

    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()